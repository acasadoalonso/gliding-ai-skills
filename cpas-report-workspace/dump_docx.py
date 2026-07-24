import sys, glob
from docx import Document
d = sorted(glob.glob(sys.argv[1] + '/**/*.docx', recursive=True))
if not d:
    sys.exit('NO DOCX FOUND in ' + sys.argv[1])
for path in d:
    doc = Document(path)
    print(f'=== {path} | paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} ===')
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
    for i, t in enumerate(doc.tables):
        print(f'--- table {i} ---')
        for r in t.rows:
            print(' | '.join(c.text for c in r.cells))
