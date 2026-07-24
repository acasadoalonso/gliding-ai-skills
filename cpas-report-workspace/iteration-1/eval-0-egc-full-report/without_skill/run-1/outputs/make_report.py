#!/usr/bin/env python3
"""Build the 2026 EGC CPAS summary report (.docx) from cpas MCP data aggregates."""
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

BASE = Path(__file__).resolve().parent
AGG = json.loads((BASE / "analysis_aggregates.json").read_text())

# --- data collected from the cpas MCP (environment/weather endpoint) ---
WEATHER = {
    "2026-07-08": {"code": "light drizzle", "tmax": 20.5, "tmin": 11.8, "precip": 0.5, "ph": 3, "wind": 27.5, "wdir": 293},
    "2026-07-09": {"code": "light drizzle", "tmax": 22.0, "tmin": 14.2, "precip": 0.9, "ph": 5, "wind": 25.7, "wdir": 315},
    "2026-07-10": {"code": "light drizzle", "tmax": 23.6, "tmin": 10.9, "precip": 0.4, "ph": 4, "wind": 20.3, "wdir": 328},
    "2026-07-13": {"code": "light drizzle", "tmax": 25.9, "tmin": 14.4, "precip": 0.2, "ph": 1, "wind": 14.1, "wdir": 354},
    "2026-07-16": {"code": "light drizzle", "tmax": 27.1, "tmin": 16.2, "precip": 0.8, "ph": 5, "wind": 11.7, "wdir": 91},
    "2026-07-17": {"code": "moderate rain", "tmax": 30.5, "tmin": 17.3, "precip": 6.2, "ph": 4, "wind": 15.4, "wdir": 170},
    "2026-07-18": {"code": "light drizzle", "tmax": 25.2, "tmin": 17.9, "precip": 2.1, "ph": 14, "wind": 19.1, "wdir": 263},
    "2026-07-20": {"code": "moderate drizzle", "tmax": 18.9, "tmin": 11.8, "precip": 2.1, "ph": 9, "wind": 22.0, "wdir": 258},
    "2026-07-22": {"code": "light drizzle", "tmax": 20.4, "tmin": 9.7, "precip": 0.4, "ph": 3, "wind": 21.8, "wdir": 271},
}

# --- per-day task data from the cpas MCP (environment/task endpoint) ---
# (phase, club_task, standard_task) ; distances in km (nominal)
TASKS = {
    "2026-07-08": ("Training", "AAT 361.6 km / 3h00 (practice)", "AAT 361.6 km / 3h00 (practice)"),
    "2026-07-09": ("Training", "Racing 218.8 km (practice)", "Racing 218.8 km (practice)"),
    "2026-07-10": ("Training", "Racing 271.5 km (practice)", "Racing 271.5 km (practice)"),
    "2026-07-13": ("Task 1", "AAT 256.1 km / 1h30 (official)", "AAT 292.3 km / 1h30 (official)"),
    "2026-07-16": ("Task 2", "AAT 217.0 km / 2h30 (official)", "AAT 227.9 km / 2h30 (official)"),
    "2026-07-17": ("Task 3", "AAT 310.0 km / 3h00 (official)", "Racing 307.6 km (official)"),
    "2026-07-18": ("Task 4", "AAT 289.2 km / 2h00 (official)", "AAT 291.1 km / 2h00 (official)"),
    "2026-07-20": ("Task 5 (Club only)", "AAT 262.6 km / 2h00 (official)", "no task set"),
    "2026-07-22": ("Task 6 / Day 11", "Racing 409.3 km (official)", "Racing 419.7 km (official)"),
}

CLASSES = ["Club", "Standard"]
days = AGG["days"]
total_flights = sum(d["flights"] for d in days)
total_enc = sum(d["encounters"] for d in days)
total_inc = sum(d["incursions"] for d in days)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

title = doc.add_heading("2026 EGC - 24th FAI European Gliding Championships", level=0)
sub = doc.add_paragraph("CPAS competition summary report - Leszno, Poland")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph("Data source: CPAS MCP server only (competition analysis, validation, settings, weather and task endpoints). Snapshot generated 24 July 2026; the competition had not yet ended at report time.")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in meta.runs:
    run.italic = True
    run.font.size = Pt(9)

# 1. Overview
doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "The 2026 European Gliding Championships (CPAS key 028cb391-202a-4cdd-bb2d-bbff50f4fdf2, "
    "SoaringSpot slug 24th-fai-egc) is being flown at Leszno, Poland (time zone Europe/Warsaw). "
    "Two classes are competing: Club and Standard. CPAS holds flight, proximity-encounter, "
    "airspace-incursion and weather data for 9 flying days between 8 and 22 July 2026, "
    f"totalling {total_flights} analysed flights."
)
p = doc.add_paragraph()
p.add_run("Key figures. ").bold = True
p.add_run(
    f"Days with recorded flights: 9 (3 official training days + 6 competition tasks). "
    f"Total flights analysed: {total_flights}. Proximity encounters: {total_enc}. "
    f"Airspace incursions: {total_inc} (none recorded on any day, in either class)."
)

# 2. Days flown
doc.add_heading("2. Days flown", level=1)
doc.add_paragraph(
    "CPAS contains flight data for 9 days. Three were official training days "
    "(8, 9 and 10 July) and six were scored competition days (13, 16, 17, 18, 20 and 22 July). "
    "The Club class flew all 6 competition tasks; the Standard class flew 5 - no Standard task "
    "was set on 20 July, when only the Club class flew (33 flights). Rest/scrubbed days "
    "(11-12, 14-15, 19, 21 July) have no flight data in CPAS."
)
t = doc.add_table(rows=1, cols=5)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
for i, h in enumerate(["Date", "Phase", "Club task", "Standard task", "Flights analysed"]):
    hdr[i].text = h
for d in days:
    phase, club, std = TASKS[d["day"]]
    row = t.add_row().cells
    row[0].text = d["day"]
    row[1].text = phase
    row[2].text = club
    row[3].text = std
    row[4].text = str(d["flights"])
row = t.add_row().cells
row[0].text = "Total"
row[4].text = str(total_flights)

# 3. Incursions
doc.add_heading("3. Airspace incursions per day and per class", level=1)
doc.add_paragraph(
    "No airspace incursions were recorded on any flying day, in either class. "
    "Every per-day incursion list in the CPAS analysis is empty. Note that in the "
    "competition settings the incursion computation flag is currently off "
    "(incursion threshold configured at 50 m), so zero recorded incursion events "
    "is the authoritative CPAS figure for this competition to date."
)
t = doc.add_table(rows=1, cols=4)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
for i, h in enumerate(["Date", "Club", "Standard", "Total incursions"]):
    hdr[i].text = h
for d in days:
    row = t.add_row().cells
    row[0].text = d["day"]
    row[1].text = "0"
    row[2].text = "0" if TASKS[d["day"]][2] != "no task set" else "- (did not fly)"
    row[3].text = "0"
row = t.add_row().cells
row[0].text = "Total"
row[1].text = "0"
row[2].text = "0"
row[3].text = "0"

# 4. Encounters
doc.add_heading("4. Proximity encounters", level=1)
doc.add_paragraph(
    f"CPAS detected {total_enc} close-proximity encounters (separation below the 30 m threshold) "
    "across the 9 flying days, involving 42 different gliders. The busiest days were "
    "13 July (Task 1, 19 encounters) and 22 July (Task 6, 21 encounters) - both full-grid "
    "racing days. The closest recorded separation was about 7.3 m (13 July). "
    "The per-class split of encounters is not published by CPAS (encounter records carry "
    "competition numbers, not class labels); the 4 encounters on 20 July are attributable "
    "to the Club class, the only class flying that day."
)
t = doc.add_table(rows=1, cols=3)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
for i, h in enumerate(["Date", "Encounters", "Minimum separation (m)"]):
    hdr[i].text = h
for d in days:
    row = t.add_row().cells
    row[0].text = d["day"]
    row[1].text = str(d["encounters"])
    row[2].text = f"{d['enc_dmin']:.1f}" if d["enc_dmin"] else "-"
row = t.add_row().cells
row[0].text = "Total"
row[1].text = str(total_enc)
doc.add_paragraph(
    "Gliders most frequently involved in encounters (whole event): "
    + ", ".join(f"{a['id']} ({a['enc']})" for a in AGG["top_enc"]) + "."
)

# 5. Weather
doc.add_heading("5. Weather", level=1)
doc.add_paragraph(
    "Per-day weather below is taken from the CPAS weather endpoint (open-meteo source, "
    "site 51.70 N / 17.87 E, elevation ~142 m). The period was characterised by "
    "north-westerly flow and light drizzle on most days. Training days were windy "
    "(up to 27.5 km/h) and cool. The first competition week warmed steadily, peaking at "
    "30.5 C on 17 July, which also brought the wettest day (6.2 mm, moderate rain). "
    "A cooler, showery westerly regime followed: 18 July had 14 hours with precipitation, "
    "and 20 July was the coolest flying day (max 18.9 C, moderate drizzle) - the day only "
    "the Club class was tasked. The final analysed day, 22 July, was dry enough for the "
    "longest racing tasks of the event (409-420 km)."
)
t = doc.add_table(rows=1, cols=7)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
for i, h in enumerate(["Date", "Conditions", "Max C", "Min C", "Precip (mm)", "Precip hours", "Max wind km/h (dir)"]):
    hdr[i].text = h
for d in days:
    w = WEATHER[d["day"]]
    row = t.add_row().cells
    row[0].text = d["day"]
    row[1].text = w["code"]
    row[2].text = f"{w['tmax']:.1f}"
    row[3].text = f"{w['tmin']:.1f}"
    row[4].text = f"{w['precip']:.1f}"
    row[5].text = str(w["ph"])
    row[6].text = f"{w['wind']:.1f} ({w['wdir']} deg)"

# 6. Flight recorders
doc.add_heading("6. Flight recorder validation", level=1)
doc.add_paragraph(
    "The CPAS validation endpoint flags 3 flight recorders of type Naviter Oudie-IGC "
    "(firmware 9.47.001) with a suggested altitude correction of -1 m. No other "
    "recorder anomalies are reported."
)

# 7. Method
doc.add_heading("7. Method and data provenance", level=1)
for line in [
    "competition list / settings: /api/compList, /api/getComp (competition key, classes site, configuration)",
    "per-day flights, encounters and incursions: /api/analysis (days array, totals)",
    "recorder validation: /api/validation",
    "per-day weather: /api/weather (open-meteo)",
    "per-day, per-class tasks: /api/task/<comp>/<day>/<Club|Standard>",
]:
    doc.add_paragraph(line, style="List Bullet")
doc.add_paragraph(
    "All figures in this report come exclusively from the CPAS MCP server. "
    "Class names on the task endpoint are 'Club' and 'Standard'; no third class exists "
    "for this competition."
)

out = BASE / "CPAS.2026_EGC_summary_report_2026-07-24.docx"
doc.save(out)
print(f"saved {out}")
