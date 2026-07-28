#!/usr/bin/env python
"""Generate the CPAS summary report (.docx) for the 2026 WGC from aggregated CPAS data."""
import json
import os
from datetime import date

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
AGG = os.path.join(HERE, "wgc2026_cpas_aggregates.json")
OUT = os.path.join(HERE, f"CPAS.wgc2026_summary_report_{date.today().isoformat()}.docx")

with open(AGG) as f:
    agg = json.load(f)

# Weather per day (from cpas environment/weather endpoint, open-meteo source)
WEATHER = {
    "2026-05-13": ("Drizzle", 12.7, 1.2, 20.2),
    "2026-05-14": ("Drizzle", 15.9, 0.1, 25.0),
    "2026-05-19": ("Drizzle", 16.9, 0.7, 8.4),
    "2026-05-20": ("Drizzle", 20.5, 0.1, 8.0),
    "2026-05-21": ("Drizzle", 19.6, 0.4, 19.7),
    "2026-05-22": ("Overcast", 22.0, 0.0, 17.0),
    "2026-05-23": ("Overcast", 22.0, 0.0, 8.9),
    "2026-05-24": ("Drizzle", 26.0, 0.7, 12.9),
    "2026-05-26": ("Clear", 25.8, 0.0, 21.1),
    "2026-05-27": ("Drizzle", 23.1, 0.5, 24.2),
    "2026-05-28": ("Overcast", 18.0, 0.0, 20.1),
    "2026-05-29": ("Overcast", 21.5, 0.0, 17.0),
}

COMP_DAYS = {"2026-05-19", "2026-05-20", "2026-05-21", "2026-05-22", "2026-05-23",
             "2026-05-24", "2026-05-26", "2026-05-27", "2026-05-28", "2026-05-29"}
TRAINING_DAYS = {"2026-05-13", "2026-05-14"}

doc = Document()

title = doc.add_heading("CPAS Summary Report — 2026 WGC", level=0)
sub = doc.add_paragraph(
    "40th FAI World Gliding Championships — Czestochowa-Rudniki, Poland\n"
    f"Report generated {date.today().isoformat()} from the CPAS analysis dataset "
    "(competition key 7a3ee794-d69c-4959-8e80-d1e3ad7ef266, folder wgc2026)."
)
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ---------------- Overview ----------------
doc.add_heading("1. Overview", level=1)
t = agg["totals"]
days = [d for d in agg["days"] if d["day"] != "practice_1"]
practice1 = next(d for d in agg["days"] if d["day"] == "practice_1")
n_comp = len([d for d in days if d["day"] in COMP_DAYS])
n_train = len([d for d in days if d["day"] in TRAINING_DAYS])

for line in [
    f"Days with recorded flying: {len(days)} calendar days "
    f"({n_train} pre-contest training days, 13-14 May, and {n_comp} competition days, "
    "19-24 and 26-29 May 2026), plus one 'practice_1' entry containing a single flight.",
    f"Total flights analysed: {t['total_flights']} (typically 73 per competition day, "
    "across the Open, 18 Metre and 20 Metre Multi-Seat classes).",
    f"Gliders tracked over the competition: {t['aircraft_total']}.",
    f"Airspace incursions recorded: {t['total_incursions']} — note that incursion detection "
    "was DISABLED in the competition settings (incursions: false) and no airspace file is "
    "loaded on the server, so this figure reflects absence of analysis, not necessarily "
    "absence of incursions.",
    f"Proximity encounters detected: {t['total_encounters']} "
    "(40 m separation threshold).",
]:
    doc.add_paragraph(line, style="List Bullet")

# ---------------- Day-by-day table ----------------
doc.add_heading("2. Day-by-day summary", level=1)
table = doc.add_table(rows=1, cols=8)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, h in enumerate(["Day", "Flights", "Encounters", "Incursions",
                       "Min sep. (m)", "Team enc.", "Weather", "Max temp / wind"]):
    hdr[i].text = h
for d in days:
    w = WEATHER.get(d["day"])
    row = table.add_row().cells
    label = d["day"]
    if d["day"] in TRAINING_DAYS:
        label += " (training)"
    row[0].text = label
    row[1].text = str(d["flights"])
    row[2].text = str(d["encounters"])
    row[3].text = str(d["incursions"])
    row[4].text = "-" if d["min_dmin"] is None else f"{d['min_dmin']:.1f}"
    row[5].text = str(d["team_encounters"])
    if w:
        row[6].text = w[0] + (f", {w[2]:.1f} mm" if w[2] else "")
        row[7].text = f"{w[1]:.0f} degC / {w[3]:.0f} km/h"
    else:
        row[6].text = row[7].text = "-"
row = table.add_row().cells
row[0].text = "practice_1"
row[1].text = str(practice1["flights"])
row[2].text = str(practice1["encounters"])
row[3].text = str(practice1["incursions"])
row[4].text = "-"
row[5].text = "0"
row[6].text = row[7].text = "-"

# ---------------- Incursions ----------------
doc.add_heading("3. Airspace incursions", level=1)
doc.add_paragraph(
    "No airspace incursions are recorded for any day or any class: the per-day incursion "
    "lists are empty on all 13 day entries (0 total). However, the competition settings show "
    "incursion detection switched off (incursions: false, threshold 50 m) and the airspace "
    "endpoint returns 'Airspace data not found', so incursion analysis was never run for "
    "this competition. The zero counts should therefore be read as 'not analysed' rather "
    "than 'no incursions occurred'. Per-class incursion counts are consequently zero for the "
    "Open, 18 Metre and 20 Metre Multi-Seat classes alike."
)

# ---------------- Encounters ----------------
doc.add_heading("4. Proximity encounters", level=1)
st = t["encounter_status"]
doc.add_paragraph(
    f"CPAS flagged {t['total_encounters']} proximity encounters over the 12 flying days "
    "(40 m distance threshold). Review status of the encounter comments:"
)
for line in [
    f"Team flying (expected close formation): {st.get('Team', 0)}",
    f"Dubious (possible data artefact): {st.get('Dubious', 0)}",
    f"Not yet reviewed / no comment: {st.get('unreviewed', 0)}",
]:
    doc.add_paragraph(line, style="List Bullet")
doc.add_paragraph(
    f"Of all encounters, {t['close_encounters_under_20m']} had a minimum separation below "
    "20 m; the closest recorded separations were about 8.1 m (22 and 26 May). "
    f"In {t['flarm_no_match_encounters']} encounters the FLARM-awareness check "
    "(geometry method, 400 m threshold) found no matching warning samples on either "
    "aircraft, i.e. the pilots may not have received a FLARM alert before the encounter. "
    "The busiest days for encounters were 20 May and 24 May (26 each), 23 May (25) and "
    "29 May (23); the quietest scored days were 28 May (8) and 21 May (10)."
)

# ---------------- Weather ----------------
doc.add_heading("5. Weather (Open-Meteo, contest site)", level=1)
doc.add_paragraph(
    "Daily weather was archived by CPAS from Open-Meteo for the contest site "
    "(50.86 N, 19.25 E, elevation 255 m). The period was mostly dry with light drizzle on "
    "several days: maximum temperatures rose from about 13 degC on the first training day "
    "to a peak of 26 degC on 24 May, with predominantly westerly to north-westerly winds "
    "(max 10 m wind speeds 8-25 km/h). Completely dry, clear-to-overcast days were "
    "22-23 May, 26 May and 28-29 May."
)

# ---------------- Flight recorders ----------------
doc.add_heading("6. Flight recorder validation notes", level=1)
doc.add_paragraph(
    "The competition-level validation endpoint returned an empty report (no validation "
    "findings stored). From the per-flight analysis data:"
)
frt = agg["fr_types"]
lx = sum(v for k, v in frt.items() if k.startswith("LXNAV"))
nav = sum(v for k, v in frt.items() if k.startswith("Naviter"))
for line in [
    f"Recorder fleet: {lx} flight logs from LXNAV devices (LX9000/9050/9070 families and "
    f"S10) and {nav} from Naviter Oudie devices, out of {t['total_flights']} logs.",
    "Barometric altitude error above 10 m was seen on four flights: "
    + ", ".join(f"CN {x['id']} ({x['baro_error']} m)" for x in agg["fr_issues"]["high_baro_error_flights"]) + ".",
    f"{agg['fr_issues']['alt_anomaly_flights_over_0_1']} flight logs showed an altitude-"
    "anomaly score above 0.1 (possible pressure-altitude noise or sensor drift); the "
    "highest recurring scores belonged to CNs UFO, OR, JP, IP3 and BR.",
    f"{agg['fr_issues']['implausible_encounter_aircraft']} aircraft records inside "
    "encounter files were marked 'plausible: false', indicating trace segments the "
    "analyser considered unreliable around those events.",
]:
    doc.add_paragraph(line, style="List Bullet")

# ---------------- Limitations ----------------
doc.add_heading("7. Data limitations", level=1)
for line in [
    "Incursion and terrain-clearance analysis were disabled in the competition settings; "
    "no airspace file is loaded, so sections on incursions report zeros by construction.",
    "Task definitions are not stored on the CPAS server for this competition (the task "
    "endpoint returns 'Task data not found'), so no per-class task tables are included.",
    "CPAS identifies flights by competition number only; class membership is not part of "
    "the CPAS dataset, so per-class breakdowns beyond the incursion statement above are "
    "not derivable from CPAS alone.",
    "The 13-14 May entries pre-date the official contest period and are treated as "
    "training days; 'practice_1' contains a single flight only.",
]:
    doc.add_paragraph(line, style="List Bullet")

for section in doc.sections:
    section.left_margin = section.right_margin = Pt(54)

doc.save(OUT)
print(OUT)
