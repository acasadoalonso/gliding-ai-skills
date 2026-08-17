#!/usr/bin/env python3
"""Build a .docx from an IGC conformance markdown report.

Usage:
    python3 make_igc_conformance_docx.py <input.md> [--out <file.docx>]

Reads the markdown report (tables, headings, blockquotes, code blocks) and
produces a formatted .docx with matching structure.
"""
import argparse
import datetime
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

REPORTS_DIR = Path("/home/angel/reports")


def parse_markdown(text: str):
    """Parse the markdown into a list of block elements."""
    blocks = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2)})
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', line.strip()):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and (lines[i].startswith(">") or (quote_lines and lines[i].strip() == "")):
                if lines[i].startswith(">"):
                    quote_lines.append(re.sub(r'^>\s?', '', lines[i]))
                else:
                    quote_lines.append("")
                i += 1
            blocks.append({"type": "blockquote", "text": "\n".join(quote_lines).strip()})
            continue

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r'^\|[\s|:-]+\|', lines[i + 1]):
            # Header
            header_cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            rows = []
            i += 2  # skip separator
            while i < len(lines) and "|" in lines[i]:
                cells = [c.strip() for c in lines[i].strip().split("|")[1:-1]]
                rows.append(cells)
                i += 1
            blocks.append({"type": "table", "headers": header_cells, "rows": rows})
            continue

        # Bullet list
        if re.match(r'^[-*]\s+', line):
            bullet_lines = []
            while i < len(lines) and re.match(r'^[-*]\s+', lines[i]):
                bullet_lines.append(re.sub(r'^[-*]\s+', '', lines[i]))
                i += 1
            blocks.append({"type": "bullet", "items": bullet_lines})
            continue

        # Bold inline markers — convert to plain text for now
        if line.strip() == "":
            i += 1
            continue

        # Paragraph (including inline bold **text**)
        para_lines = []
        while i < len(lines) and lines[i].strip() != "" and not re.match(r'^#{1,6}\s', lines[i]) and not lines[i].startswith("|") and not lines[i].startswith(">") and not lines[i].strip().startswith("```") and not re.match(r'^[-*]\s+', lines[i]):
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            blocks.append({"type": "paragraph", "text": "\n".join(para_lines)})

        i += 1
    return blocks


def bold_run(run, text: str):
    """Make a run bold."""
    run.font.bold = True
    run.text = text


def add_inline(p, text: str):
    """Add text to a paragraph, rendering **bold**, *italic* and `code` spans."""
    for part in re.split(r'(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)', text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).font.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*"):
            p.add_run(part[1:-1]).font.italic = True
        else:
            p.add_run(part)


def add_table(doc: Document, headers: list, rows: list):
    """Add a styled table with bold header row."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for r in t.rows[0].cells[i].paragraphs:
            for run in r.runs:
                run.font.bold = True
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            add_inline(cell.paragraphs[0], cell_text)


def build_docx(blocks: list) -> Document:
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for block in blocks:
        if block["type"] == "heading":
            # Heading styles carry their own formatting; drop the inline markers.
            doc.add_heading(re.sub(r'[*`]', '', block["text"]), level=block["level"])

        elif block["type"] == "hr":
            doc.add_paragraph("─" * 40)

        elif block["type"] == "blockquote":
            p = doc.add_paragraph(block["text"])
            p.italic = True
            for run in p.runs:
                run.font.size = Pt(10)

        elif block["type"] == "code":
            p = doc.add_paragraph(block["text"])
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)

        elif block["type"] == "table":
            add_table(doc, block["headers"], block["rows"])

        elif block["type"] == "bullet":
            for item in block["items"]:
                add_inline(doc.add_paragraph(style="List Bullet"), item)

        elif block["type"] == "paragraph":
            add_inline(doc.add_paragraph(), block["text"])

    return doc


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input", type=Path, help="Path to the .md report")
    ap.add_argument("--out", type=Path, default=None)
    a = ap.parse_args()

    if not a.input.exists():
        print(f"ERROR: {a.input} not found", file=sys.stderr)
        sys.exit(1)

    text = a.input.read_text(encoding="utf-8")
    blocks = parse_markdown(text)
    doc = build_docx(blocks)

    out = a.out or REPORTS_DIR / (
        a.input.stem + ".docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"saved {out}")


if __name__ == "__main__":
    main()
